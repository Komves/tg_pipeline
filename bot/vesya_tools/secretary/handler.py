from __future__ import annotations

import re
import os
import imaplib
import email
import tempfile
import base64
import quopri
import json
import html as html_lib
from pathlib import Path
from email.header import decode_header
from email.utils import parsedate_to_datetime
from datetime import datetime, timedelta
import openai
from aiogram.types import InlineKeyboardMarkup, InlineKeyboardButton
from docx import Document 
import copy



def _compress_emails(emails):
    compressed = []

    for e in emails:
        compressed.append({
            "from": e.get("from"),
            "to": e.get("to"),
            "folder": e.get("folder"),
            "subject": e.get("subject"),
            "date": e.get("date"),
            "attachments": e.get("attachments", [])[:2],
            "body": (e.get("body") or "")[:1200],
        })

    return compressed

SECRETARY_CACHE = {}

# =========================
# EMAIL UTILS
# =========================

def _handle_email_selection(message, text, cache):
    nums = [x.strip() for x in text.split(",")]

    selected = []

    for n in nums:
        if not n.isdigit():
            continue

        idx = int(n) - 1
        if 0 <= idx < len(cache["emails"]):
            selected.append(cache["emails"][idx])

    cache["selected"] = selected
    return selected


def _decode_mime(text):
    if not text:
        return ""

    parts = decode_header(text)
    out = ""

    for part, enc in parts:
        if isinstance(part, bytes):
            try:
                out += part.decode(enc or "utf-8", errors="ignore")
            except Exception:
                out += part.decode("utf-8", errors="ignore")
        else:
            out += str(part)

    return out




def _imap_connect():
    user = os.getenv("MAILRU_EMAIL", "comves@list.ru")
    password = os.getenv("MAILRU_APP_PASSWORD")

    if not password:
        return None

    imap = imaplib.IMAP4_SSL("imap.mail.ru", 993)
    imap.login(user, password)
    return imap

def _imap_select_mailbox(imap, mailbox):
    mailbox = str(mailbox or "INBOX")

    for candidate in (mailbox, f'"{mailbox}"'):
        try:
            status, _ = imap.select(candidate)
            if status == "OK":
                return True
        except Exception:
            continue

    return False


def _parse_imap_list_name(raw_line):
    if isinstance(raw_line, bytes):
        line = raw_line.decode("utf-8", errors="ignore")
    else:
        line = str(raw_line)

    m = re.search(r'\) "[^"]*" (.+)$', line)
    if not m:
        return ""

    name = m.group(1).strip()

    if name.startswith('"') and name.endswith('"'):
        name = name[1:-1]

    return name.strip()


def _mailru_target_folders(imap):
    folders = [
        {
            "mailbox": "INBOX",
            "folder": "INBOX",
        }
    ]

    try:
        status, data = imap.list()

        if status == "OK":
            for raw in data or []:
                line = raw.decode("utf-8", errors="ignore") if isinstance(raw, bytes) else str(raw)
                name = _parse_imap_list_name(line)

                if not name:
                    continue

                low = line.lower()

                print(f"[secretary] mailbox found: {line}", flush=True)

                if (
                    "\\sent" in low
                    or "sent" in low
                    or "отправ" in low
                    or "&bb4" in low
                ):
                    folders.append({
                        "mailbox": name,
                        "folder": "Sent",
                    })

    except Exception as e:
        print(f"[secretary] mailbox list failed: {type(e).__name__}: {e}", flush=True)

    folders.append({
        "mailbox": "Sent",
        "folder": "Sent",
    })

    unique = []
    seen = set()

    for f in folders:
        key = f"{f.get('mailbox')}|{f.get('folder')}"
        if key not in seen:
            unique.append(f)
            seen.add(key)

    return unique

def _parse_email_date(value):
    try:
        dt = parsedate_to_datetime(value or "")
        if dt is None:
            return None
        if dt.tzinfo is not None:
            dt = dt.astimezone().replace(tzinfo=None)
        return dt
    except Exception:
        return None


def _fetch_mailru_headers(limit=300, period_start=None, period_end=None):
    imap = _imap_connect()
    if imap is None:
        return []

    results = []

    start_dt = period_start
    end_dt = period_end

    if end_dt:
        end_dt = end_dt + timedelta(days=1)

    try:
        for folder_info in _mailru_target_folders(imap):
            try:
                mailbox = folder_info["mailbox"]
                folder = folder_info["folder"]

                if not _imap_select_mailbox(imap, mailbox):
                    print(f"[secretary] cannot select mailbox={mailbox}", flush=True)
                    continue
                status, data = imap.search(None, "ALL")

                if status != "OK" or not data or not data[0]:
                    continue

                ids = data[0].split()[-limit:]

                for num in ids:
                    status, msg_data = imap.fetch(
                        num,
                        "(BODY.PEEK[HEADER.FIELDS (DATE FROM TO SUBJECT MESSAGE-ID IN-REPLY-TO REFERENCES)])"
                    )

                    if status != "OK" or not msg_data or not msg_data[0]:
                        continue

                    raw_headers = msg_data[0][1]
                    msg = email.message_from_bytes(raw_headers)

                    date_raw = _decode_mime(msg.get("Date"))
                    msg_dt = _parse_email_date(date_raw)

                    if start_dt and msg_dt and msg_dt < start_dt:
                        continue

                    if end_dt and msg_dt and msg_dt >= end_dt:
                        continue

                    results.append({
                        "folder": folder,
                        "imap_folder": mailbox,
                        "imap_id": num.decode(errors="ignore"),
                        "date": date_raw,
                        "date_iso": msg_dt.isoformat() if msg_dt else "",
                        "from": _decode_mime(msg.get("From")),
                        "to": _decode_mime(msg.get("To")),
                        "subject": _decode_mime(msg.get("Subject")),
                        "message_id": _decode_mime(msg.get("Message-ID")),
                        "in_reply_to": _decode_mime(msg.get("In-Reply-To")),
                        "references": _decode_mime(msg.get("References")),
                        "attachments": [],
                        "body": "",
                    })

            except Exception as e:
                print(f"[secretary] header fetch folder failed {folder}: {type(e).__name__}: {e}", flush=True)
                continue

    finally:
        try:
            imap.logout()
        except Exception:
            pass

    return results


def _fetch_selected_full_emails(headers):
    imap = _imap_connect()
    if imap is None:
        return headers

    out = []

    try:
        for h in headers:
            try:
                mailbox = h.get("imap_folder") or "INBOX"
                imap_id = h.get("imap_id")

                if not imap_id:
                    continue

                if not _imap_select_mailbox(imap, mailbox):
                    continue

                status, msg_data = imap.fetch(
                    str(imap_id).encode(),
                    "(RFC822)"
                )

                if status != "OK" or not msg_data or not msg_data[0]:
                    continue

                raw = _extract_fetch_bytes(msg_data)
                if not raw:
                    continue

                msg = email.message_from_bytes(raw)
                body, attachments = _extract_light_message(msg, str(imap_id))

                item = dict(h)
                item["body"] = body
                item["attachments"] = attachments
                out.append(item)

            except Exception as e:
                print(
                    f"[secretary] full fetch failed imap_id={h.get('imap_id')}: {type(e).__name__}: {e}",
                    flush=True,
                )
                out.append(h)

    finally:
        try:
            imap.logout()
        except Exception:
            pass

    return out

def _imap_unquote(value):
    if value is None:
        return ""

    if isinstance(value, bytes):
        value = value.decode("utf-8", errors="ignore")

    value = str(value).strip()

    if value.startswith('"') and value.endswith('"'):
        value = value[1:-1]

    return value


def _decode_transfer_payload(raw_bytes, encoding):
    encoding = (encoding or "").lower().strip()

    if not raw_bytes:
        return b""

    if encoding == "base64":
        return base64.b64decode(raw_bytes)

    if encoding in ("quoted-printable", "quotedprintable"):
        return quopri.decodestring(raw_bytes)

    return raw_bytes


def _extract_fetch_bytes(msg_data):
    for item in msg_data or []:
        if isinstance(item, tuple) and len(item) >= 2 and isinstance(item[1], (bytes, bytearray)):
            return bytes(item[1])
    return b""




def _find_param_value(params, name):
    if not isinstance(params, list):
        return ""

    name = (name or "").lower()

    for i in range(0, len(params) - 1, 2):
        if str(params[i]).lower() == name:
            return _decode_mime(_imap_unquote(params[i + 1]))

    return ""

def _extract_light_message(msg, imap_id: str):
    """
    УПРОЩЁННОЕ ИЗВЛЕЧЕНИЕ ПИСЬМА (ДЛЯ АКТА)
    """

    body = ""
    attachments = []

    for part in msg.walk():
        content_type = (part.get_content_type() or "").lower()
        filename = part.get_filename()
        payload = part.get_payload(decode=True) or b""

        # BODY
        if content_type == "text/plain" and not filename and not body:
            try:
                body = payload.decode("utf-8", errors="ignore")
            except Exception:
                body = ""

        # ATTACHMENTS (ТОЛЬКО МЕТА)
        if filename:
            attachments.append({
                "filename": filename,
                "size": len(payload),
                "type": content_type
            })

    # ЖЁСТКИЕ ЛИМИТЫ (ВАЖНО)
    body = (body or "")[:3000]
    attachments = attachments[:3]

    return body, attachments



def _fetch_part_bytes(imap, imap_id, part_no):
    status, data = imap.fetch(
        str(imap_id).encode(),
        f"(BODY.PEEK[{part_no}])"
    )

    if status != "OK":
        return b""

    return _extract_fetch_bytes(data)


def _html_to_text(value):
    text = value or ""
    text = re.sub(r"(?is)<(script|style).*?>.*?</\1>", " ", text)
    text = re.sub(r"(?is)<br\s*/?>", "\n", text)
    text = re.sub(r"(?is)</p\s*>", "\n", text)
    text = re.sub(r"(?is)<.*?>", " ", text)
    text = html_lib.unescape(text)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n\s*\n+", "\n\n", text)
    return text.strip()

def _build_cases(emails):
    cases = {}

    for e in emails:
        sender = (e.get("from") or "unknown").lower().strip()

        if sender not in cases:
            cases[sender] = {
                "emails": []
            }

        cases[sender]["emails"].append({
            "subject": e.get("subject"),
            "body": (e.get("body") or "")[:1500],
            "attachments": e.get("attachments", []),
            "date": e.get("date")
        })

    return list(cases.values())


def _build_chain(selected_emails, actor):
    return [
        e for e in selected_emails
        if actor.lower() in (e.get("from") or "").lower()
    ]


def _gpt_analyze_chain(chain, actor):
    text = ""

    for e in chain:
        imap_id = e.get("imap_id", "")
        message_id = e.get("message_id", "")

        text += f"""
        EMAIL #{i} | IMAP_ID={imap_id} | MSG_ID={message_id}
        FOLDER: {e.get('folder')}
        DATE: {e.get('date')}
        FROM: {e.get('from')}
        TO: {e.get('to')}
        SUBJECT: {e.get('subject')}
        ATTACHMENTS: {", ".join([
            (
                a.get("filename", "")
                + (f" ({a.get('type', '')}, size={a.get('size', '')})" if a.get("size") else "")
            )
            for a in (e.get("attachments") or [])
        ])}
        BODY: {e.get('body','')[:2000]}
        ----------------------
        """

    client = openai.OpenAI()

    resp = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[
            {
                "role": "system",
                "content": "Ты юридический аналитик. Возвращай строго структурированный анализ."
            },
            {
                "role": "user",
                "content": f"""
Проанализируй переписку по участнику: {actor}

Верни JSON:
{{
  "summary": "",
  "facts": [],
  "risks": [],
  "obligations": []
}}

Переписка:
{text}
"""
            }
        ]
    )

    return resp.choices[0].message.content


# =========================
# DOCX ENGINE
# =========================

MONTH_NAMES_RU = {
    1: "Январь",
    2: "Февраль",
    3: "Март",
    4: "Апрель",
    5: "Май",
    6: "Июнь",
    7: "Июль",
    8: "Август",
    9: "Сентябрь",
    10: "Октябрь",
    11: "Ноябрь",
    12: "Декабрь",
}


def _calendar_keyboard(kind: str, year: int | None = None, month: int | None = None):
    now = datetime.now()

    year = int(year or now.year)
    month = int(month or now.month)

    first = datetime(year, month, 1)

    if month == 12:
        next_month = datetime(year + 1, 1, 1)
    else:
        next_month = datetime(year, month + 1, 1)

    days_count = (next_month - first).days

    prev_month = month - 1
    prev_year = year
    if prev_month < 1:
        prev_month = 12
        prev_year -= 1

    next_month_num = month + 1
    next_year = year
    if next_month_num > 12:
        next_month_num = 1
        next_year += 1

    rows = []

    rows.append([
        InlineKeyboardButton(text="◀", callback_data=f"sec:cal:{kind}:{prev_year}:{prev_month}"),
        InlineKeyboardButton(text=f"{MONTH_NAMES_RU.get(month, month)} {year}", callback_data="sec:none"),
        InlineKeyboardButton(text="▶", callback_data=f"sec:cal:{kind}:{next_year}:{next_month_num}"),
    ])

    rows.append([
        InlineKeyboardButton(text="Пн", callback_data="sec:none"),
        InlineKeyboardButton(text="Вт", callback_data="sec:none"),
        InlineKeyboardButton(text="Ср", callback_data="sec:none"),
        InlineKeyboardButton(text="Чт", callback_data="sec:none"),
        InlineKeyboardButton(text="Пт", callback_data="sec:none"),
        InlineKeyboardButton(text="Сб", callback_data="sec:none"),
        InlineKeyboardButton(text="Вс", callback_data="sec:none"),
    ])

    week = []
    first_weekday = first.weekday()

    for _ in range(first_weekday):
        week.append(InlineKeyboardButton(text=" ", callback_data="sec:none"))

    for day in range(1, days_count + 1):
        week.append(
            InlineKeyboardButton(
                text=str(day),
                callback_data=f"sec:date:{kind}:{year}-{month:02d}-{day:02d}",
            )
        )

        if len(week) == 7:
            rows.append(week)
            week = []

    if week:
        while len(week) < 7:
            week.append(InlineKeyboardButton(text=" ", callback_data="sec:none"))
        rows.append(week)

    return InlineKeyboardMarkup(inline_keyboard=rows)


def _actor_key(raw_from: str) -> str:
    raw = raw_from or "unknown"
    match = re.search(r'[\w\.-]+@[\w\.-]+', raw)
    return match.group(0).lower() if match else raw.lower().strip()

def _norm_msg_id(value: str) -> str:
    return (value or "").strip().lower().strip("<>")


def _header_refs(header: dict) -> set[str]:
    refs = set()

    for field in ("message_id", "in_reply_to", "references"):
        raw = str(header.get(field) or "")
        for item in re.findall(r"<([^>]+)>", raw):
            item = _norm_msg_id(item)
            if item:
                refs.add(item)

    return refs


def _expand_headers_by_threads(all_headers, seed_headers):
    seed_actors = frozenset(
        _actor_key(e.get("from", "")) for e in seed_headers
    )

    seed_senders = seed_actors
    selected = list(seed_headers or [])

    seed_thread_ids = set()
    for h in seed_headers:
        seed_thread_ids.update(_header_refs(h))

    seed_ids_flat = set(seed_thread_ids)

    selected_ids = set(seed_ids_flat)

    for h in selected:
        selected_ids.update(_header_refs(h))

    changed = True

    while changed:
        changed = False

        for h in all_headers:
            refs = _header_refs(h)

            sender = _actor_key(h.get("from", ""))

            ref_match = bool(refs and refs.intersection(selected_ids))
            actor_match = sender in seed_senders

            if not (ref_match or actor_match):
                continue

            if h in selected:
                continue

            selected.append(h)
            selected_ids.update(refs)
                
            changed = True

    return selected


def _norm_subject(value: str) -> str:
    s = (value or "").lower().strip()
    s = re.sub(r"^\s*(re|fw|fwd|ответ|пересл):\s*", "", s, flags=re.I)
    s = re.sub(r"\s+", " ", s)
    return s.strip()


def _email_dt(e: dict) -> datetime:
    iso = e.get("date_iso") or ""
    try:
        return datetime.fromisoformat(iso)
    except Exception:
        return datetime.min


def _email_participants(e: dict) -> set[str]:
    out = set()

    for field in ("from", "to"):
        raw = e.get(field) or ""
        for m in re.findall(r'[\w\.-]+@[\w\.-]+', raw):
            out.add(m.lower())

    return out


def _build_mail_threads(emails):
    emails = copy.deepcopy(emails)

    """
    УПРОЩЁННАЯ ВЕРСИЯ:
    без Message-ID графа, без union-find, без сложных связок
    """

    groups = {}

    for i, e in enumerate(emails, start=1):

        e = dict(e)
        e["email_no"] = i

        subject = (e.get("subject") or "").lower().strip()
        subject = re.sub(r"^(re:|fw:|fwd:)\s*", "", subject)
        subject = re.sub(r"\s+", " ", subject)

        sender = (e.get("from") or "").lower().strip()

        # ключ = отправитель + тема (простая группировка)
        refs = _header_refs(e)

        if refs:
            key = f"msg::{sorted(list(refs))[0]}"
        else:
            key = f"{sender}||{subject}"

        if key not in groups:
            groups[key] = {
                "thread_no": len(groups) + 1,
                "subject": subject,
                "participants": set(),
                "messages": []
            }

        groups[key]["messages"].append(e)
        groups[key]["participants"].add(sender)

    threads = list(groups.values())

    for t in threads:
        t["participants"] = list(t["participants"])

        # сортировка по дате
        try:
            t["messages"].sort(key=lambda x: x.get("date_iso") or "")
        except Exception:
            pass

    return threads

def _filter_noise_emails(emails):
    """
    Убирает мусор ДО GPT:
    - уведомления
    - системные письма
    - пустые авто-алерты
    - сервисные спамы
    """

    BAD_KEYWORDS = [
        "no-reply",
        "noreply",
        "notification",
        "уведомлен",
        "alert",
        "system",
        "mailer-daemon",
        "do not reply",
        "password",
        "код подтверждения",
        "verification",
        "security",
        "robot",
        "автоматичес"
    ]

    def is_noise(e):
        text = " ".join([
            str(e.get("subject", "")),
            str(e.get("from", "")),
            str(e.get("body", "")[:300])
        ]).lower()

        # 1. пустые письма
        if not e.get("attachments") and not e.get("subject"):
            return True

        # 2. слишком короткие авто-письма
        if len(e.get("subject") or "") < 2 and not e.get("attachments"):
            return True

        # 3. ключевые слова мусора
        if any(k in text for k in BAD_KEYWORDS):
            return True

        return False

    clean = [e for e in emails if not is_noise(e)]

    return clean

def _thread_to_text(thread):
    def clean_body(text):
        text = text or ""

        # убираем мусор цитат
        text = re.sub(r">.*", "", text)
        text = re.sub(r"-----Original Message-----.*", "", text, flags=re.S)

        # схлопываем пробелы
        text = re.sub(r"\s+", " ", text).strip()

        return text

    out = []

    for e in thread.get("messages") or []:

        body = clean_body(e.get("body"))

        # ❗ ВАЖНО: оставляем только первые 2-3 смысловые строки
        # режем не по символам, а по информации
        body = body[:500]

        attachments = e.get("attachments") or []

        # фильтр "пустых писем"
        if not body and not attachments:
            continue

        out.append(
            f"""
FROM: {e.get('from')}
SUBJECT: {e.get('subject')}
FOLDER: {e.get('folder')}
ATT: {len(attachments)}
BODY: {body}
"""
        )

    result = "\n---\n".join(out)

    # жесткий потолок, но уже после очистки
    return result[:9000]

def _safe_json_loads(raw):
    raw = (raw or "").strip()

    raw = re.sub(r"^```json\s*", "", raw, flags=re.I)
    raw = re.sub(r"^```\s*", "", raw)
    raw = re.sub(r"\s*```$", "", raw)

    return json.loads(raw)


def _gpt_analyze_thread(thread, project_name, period_text):
    client = openai.OpenAI()

    text = _thread_to_text(thread)

    resp = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[
            {
                "role": "system",
                "content": (
                    "Ты анализируешь ОДНУ цепочку деловой переписки.\n"
                    "Не пересказывай письма.\n"
                    "Определи процесс: запрос → действия → результат.\n"
                    "INBOX = запрос\n"
                    "Sent = действие пользователя (важнее INBOX)\n"
                )
            },
            {
                "role": "user",
                "content": f"""
Проект: {project_name}
Период: {period_text}

Верни JSON:
{{
  "thread_no": {thread.get("thread_no")},
  "task": "",
  "request_from": "",
  "request": "",
  "action": "",
  "result": "",
  "status": "в работе",
  "emails": []
}}

Цепочка:
{text}
"""
            }
        ]
    )

    raw = resp.choices[0].message.content.strip()

    raw = re.sub(r"^```json", "", raw)
    raw = re.sub(r"^```", "", raw)
    raw = re.sub(r"```$", "", raw)

    try:
        return json.loads(raw)
    except Exception:
        return {
            "thread_no": thread.get("thread_no"),
            "task": "parse_error",
            "raw": raw
        }

def _gpt_make_tasks_from_threads(thread_summaries, project_name, period_text):
    client = openai.OpenAI()

    thread_summaries = thread_summaries[:20]  # ЖЁСТКИЙ ЛИМИТ

    resp = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[
            {
                "role": "system",
                "content": (
                    "Ты формируешь акт выполненных работ.\n"
                    "ЗАПРЕЩЕНО:\n"
                    "- использовать нумерацию писем (EMAIL #, №, списки с цифрами)\n"
                    "- выводить номера писем в виде 1,2,3...\n"
                    "- использовать IMAP как список через запятую без пояснения\n"
                    "\n"
                    "Вместо этого:\n"
                    "- описывай письма словами (тема, суть)\n"
                    "- если нужны ссылки на письма — используй IMAP_ID только внутри JSON массива, не в тексте\n"
                    "Каждый thread = одна задача.\n"
                    "Не дроби и не объединяй разные threads.\n"
                )
            },
            {
                "role": "user",
                "content": f"""
Проект: {project_name}
Период: {period_text}

Threads:
{json.dumps(thread_summaries, ensure_ascii=False, indent=2)}

Верни JSON массив:
[
  {{
    "request_from": "",
    "topic": "",
    "done": "",
    "status": "в работе"
  }}
]
"""
            }
        ]
    )

    raw = resp.choices[0].message.content.strip()

    raw = re.sub(r"^```json", "", raw)
    raw = re.sub(r"^```", "", raw)
    raw = re.sub(r"```$", "", raw)

    try:
        return json.loads(raw)
    except Exception:
        return [
            {
                "task": "parse_error",
                "done": raw,
                "status": "error"
            }
        ]


def _collect_actors(emails):
    actors = {}

    for e in emails:
        key = _actor_key(e.get("from", ""))
        if key not in actors:
            actors[key] = {
                "key": key,
                "name": e.get("from", key),
                "count": 0
            }
        actors[key]["count"] += 1

    return list(actors.values())


def _tasks_review_text(tasks):
    lines = ["GPT сгруппировал задачи. Проверь статусы:\n"]
    

    for task in tasks:
        i = task.get("task_id", 0)
        emails = task.get("emails") or []
        emails_text = "\n".join(
            f"• {x.get('subject','(без темы)')} — {x.get('from','')} — {x.get('date','')}"
            for x in emails
            if isinstance(x, dict)
        )

        request_from = task.get("request_from") or task.get("from") or task.get("initiator") or ""
        topic = task.get("topic") or task.get("task") or task.get("subject") or "Тема не определена"
        done = task.get("done") or task.get("action") or task.get("result") or "Описание не сформировано"
        status = task.get("status") or "требует проверки"   

        lines.append(
                f"{i}. Получен запрос от: {request_from if request_from else 'не определено'}\n"
            f"Тема: {topic}\n"
            f"Что сделано: {done}\n"
            f"Статус задачи: {status}\n"
        )

    lines.append("Когда статусы верные — нажми «Сформировать акт».")
    return "\n".join(lines)


def _tasks_review_keyboard(tasks):
    rows = []

    for task in tasks:
        task_id = int(task.get("task_id") or 0)

        rows.append([
            InlineKeyboardButton(text=f"{task_id}: закрыто", callback_data=f"sec:status:{task_id}:закрыто"),
            InlineKeyboardButton(text=f"{task_id}: в работе", callback_data=f"sec:status:{task_id}:в работе"),
        ])
        rows.append([
            InlineKeyboardButton(text=f"{task_id}: убрать", callback_data=f"sec:status:{task_id}:убрать"),
        ])

    rows.append([
        InlineKeyboardButton(text="Сформировать акт", callback_data="sec:act:make")
    ])

    return InlineKeyboardMarkup(inline_keyboard=rows)

def _make_docx(tasks):
    doc = Document()

    doc.add_heading("Акт выполненных работ", level=1)

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"

    hdr = table.rows[0].cells
    hdr[0].text = "№"
    hdr[1].text = "Задача / вопрос"
    hdr[2].text = "Что сделано"
    hdr[3].text = "Статус"

    for i, task in enumerate(tasks, start=1):
        row = table.add_row().cells
        row[0].text = str(i)
        row[1].text = str(task.get("task", ""))
        row[2].text = str(task.get("done", ""))
        row[3].text = str(task.get("status", ""))

    path = "/tmp/secretary_act.docx"
    doc.save(path)

    return path


def _gpt_make_tasks(emails, project_name, period_text):
    text = ""

    for e in emails:
        imap_id = e.get("imap_id", "")
        subject = e.get("subject", "")
        sender = e.get("from", "")
        date = e.get("date", "")

        text += f"""
📩 {subject}
👤 {sender}
📅 {date}
🆔 IMAP_ID: {imap_id}

📝 {e.get('body','')[:2000]}

-------------------
"""
    

    client = openai.OpenAI()

    resp = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[
    {
        "role": "system",
        "content": (
            "Ты формируешь АКТ ВЫПОЛНЕННЫХ КОНСУЛЬТАЦИОННЫХ УСЛУГ.\n"
            "\n"
            "КРИТИЧЕСКОЕ ОГРАНИЧЕНИЕ:\n"
            "- НЕ объединяй письма в кейсы\n"
            "- НЕ дроби письма\n"
            "- НЕ создавай логические цепочки между письмами\n"
            "- НЕ интерпретируй переписку как единый процесс\n"
            "\n"
            "ПРИНЦИП РАБОТЫ:\n"
            "- каждое письмо — это отдельный входной факт\n"
            "- ты НЕ определяешь структуру переписки\n"
            "- ты НЕ группируешь письма\n"
            "\n"
            "ТВОЯ РОЛЬ:\n"
            "- понять смысл каждого письма\n"
            "- описать результат консультационной работы\n"
            "- оформить деловым юридическим языком\n"
            "\n"
            "ЗАПРЕЩЕНО:\n"
            "- склейка писем\n"
            "- восстановление кейсов\n"
            "- построение цепочек событий\n"
        )
    },
    {
        "role": "user",
        "content": f"""
Проект/группа: {project_name}
Период: {period_text}

Верни СТРОГО JSON-массив. Без markdown. Без пояснений.

Правила формирования задач:
...
"""
        }
    ]
)

    raw = resp.choices[0].message.content.strip()

    with open("/tmp/gpt_answer.txt", "w", encoding="utf-8") as f:
        f.write(raw)

    print("GPT ANSWER SAVED: /tmp/gpt_answer.txt", flush=True)

    try:
        import json
        data = json.loads(raw)
        if isinstance(data, list):
            return data
    except Exception:
        pass

    return [
        {
            "task": "Анализ переписки",
            "done": raw,
            "status": "требует проверки"
        }
    ]

async def handle_secretary_callback(cb) -> bool:
    data = cb.data or ""

    if data == "sec:none":
        await cb.answer()
        return True

    chat_id = cb.message.chat.id
    cache = SECRETARY_CACHE.get(chat_id)

    if cache is None:
        await cb.answer("Сценарий секретаря не запущен")
        return True

    if data.startswith("sec:status:"):
        parts = data.split(":")
        idx = int(parts[2])
        status = parts[3]

        tasks = cache.get("tasks") or []

        if 0 <= idx < len(tasks):
            if status == "убрать":
                tasks[idx]["exclude"] = True
                tasks[idx]["status"] = "убрать из акта"
            else:
                tasks[idx]["exclude"] = False
                tasks[idx]["status"] = status

            cache["tasks"] = tasks

            try:
                await cb.message.edit_text(
                    _tasks_review_text(tasks),
                    reply_markup=_tasks_review_keyboard(tasks),
                )
            except Exception as e:
                if "message is not modified" not in str(e).lower():
                    raise

        await cb.answer()
        return True

    if data == "sec:act:make":
        tasks = [
            t for t in (cache.get("tasks") or [])
            if not t.get("exclude")
        ]

        file_path = _make_docx(tasks)

        cache["stage"] = "done"
        cache["doc"] = file_path

        from aiogram.types import FSInputFile

        await cb.message.answer("Готово. Акт сформирован.")

        await cb.message.answer_document(
            FSInputFile(file_path),
            caption="Проект акта выполненных работ"
        )

        await cb.answer()
        return True


    if data.startswith("sec:cal:"):
        parts = data.split(":")
        kind = parts[2]
        year = int(parts[3])
        month = int(parts[4])

        title = "Выбери дату начала периода:" if kind == "start" else "Выбери дату окончания периода:"

        await cb.message.edit_text(
            title,
            reply_markup=_calendar_keyboard(kind, year, month),
        )
        await cb.answer()
        return True

    if data.startswith("sec:date:"):
        parts = data.split(":")
        kind = parts[2]
        date_value = parts[3]

        selected_date = datetime.strptime(date_value, "%Y-%m-%d")

        if kind == "start":
            cache["period_start"] = selected_date
            cache["stage"] = "wait_period_end"

            await cb.message.edit_text(
                f"Дата начала: {selected_date.strftime('%d.%m.%Y')}\n\n"
                "Теперь выбери дату окончания периода:",
                reply_markup=_calendar_keyboard("end", selected_date.year, selected_date.month),
            )
            await cb.answer()
            return True

        if kind == "end":
            start = cache.get("period_start")
            if not start:
                await cb.answer("Сначала выбери дату начала")
                return True

            if selected_date < start:
                await cb.answer("Дата окончания раньше даты начала")
                return True

            cache["period_end"] = selected_date
            cache["period_text"] = f"{start.strftime('%d.%m.%Y')}-{selected_date.strftime('%d.%m.%Y')}"
            cache["stage"] = "actor_select"

            await cb.message.edit_text(
                f"Период выбран: {cache['period_text']}\n\n"
                "Читаю почту и собираю адресатов..."
            )

            emails = _fetch_mailru_headers(
                limit=300,
                period_start=cache.get("period_start"),
                period_end=cache.get("period_end"),
            )
            cache["emails"] = emails

            actors = _collect_actors(emails)
            cache["actors"] = actors

            if not actors:
                await cb.message.answer("Писем в ящике не нашла.")
                cache["stage"] = "idle"
                await cb.answer()
                return True

            await cb.message.answer(
                "Выбери адресатов, которые относятся к проекту.\n"
                "Пока без кнопок: напиши номера через запятую.\n\n" +
                "\n".join([
                    f"{i+1}. ☐ {a['name']} ({a['count']})"
                    for i, a in enumerate(actors)
                ])
            )

            await cb.answer()
            return True

    await cb.answer()
    return True

def _classify_domain(e):
    text = f"{e.get('subject','')} {e.get('body','')}".lower()

    if "резюме" in text or "кадры" in text or "отдел кадров" in text:
        return "hr"

    if "койко" in text or "общежит" in text or "регистрац" in text:
        return "housing"

    if "прокурат" in text or "момвд" in text or "запрос" in text:
        return "gov"

    if "суд" in text or "дело" in text or "а40" in text:
        return "court"

    return "legal"


async def handle_secretary_message(message, text: str, object_text=None) -> bool:

    if not isinstance(text, str):
        text = str(text)

    raw = text.strip()
    t = raw.lower()

    chat_id = message.chat.id

    cache = SECRETARY_CACHE.get(chat_id)

    if cache is None:
        cache = SECRETARY_CACHE[chat_id] = {
            "stage": "idle",
            "project": "",
            "period_text": "",
            "period_start": None,
            "period_end": None,
            "emails": [],
            "actors": [],
            "selected_actors": [],
            "selected": [],
            "tasks": []
        }

    # -----------------------------
    # START ACT SCENARIO
    # -----------------------------
    m = re.search(r"составь\s+акт\s+по\s+(.+)$", t, flags=re.I)
    if m:
        project = (m.group(1) or "").strip()
        project = project.upper() if project else "ПРОЕКТ"

        cache.clear()
        cache.update({
            "stage": "wait_period_start",
            "project": project,
            "period_text": "",
            "period_start": None,
            "period_end": None,
            "emails": [],
            "actors": [],
            "selected_actors": [],
            "selected": [],
            "tasks": []
        })

        now = datetime.now()

        await message.answer(
            f"Ок. Составляем акт по {project}.\n\n"
            "Выбери дату начала периода:",
            reply_markup=_calendar_keyboard("start", now.year, now.month),
        )
        return True

    # -----------------------------
    # PERIOD INPUT
    # -----------------------------
    if False and cache.get("stage") == "wait_period":
        period = _parse_period(raw)

        if not period:
            await message.answer(
                "Не поняла период. Выбери кнопку или напиши так:\n"
                "01.06.2026-30.06.2026",
                reply_markup=_period_keyboard(),
            )
            return True

        start, end, period_text = period

        cache["period_start"] = start
        cache["period_end"] = end
        cache["period_text"] = period_text
        cache["stage"] = "actor_select"

        emails = _fetch_mailru_headers(limit=50)
        cache["emails"] = emails

        actors = _collect_actors(emails)
        cache["actors"] = actors

        if not actors:
            await message.answer("Писем в ящике не нашла.")
            cache["stage"] = "idle"
            return True

        await message.answer(
            "Выбери адресатов, которые относятся к проекту.\n"
            "Пока без кнопок: напиши номера через запятую.\n\n" +
            "\n".join([
                f"{i+1}. ☐ {a['name']} ({a['count']})"
                for i, a in enumerate(actors)
            ])
        )
        return True

    # -----------------------------
    # ACTOR SELECTION
    # -----------------------------
    if cache.get("stage") == "actor_select":

        nums = [int(x.strip()) for x in t.split(",") if x.strip().isdigit()]

        selected_actors = []
        for n in nums:
            idx = n - 1
            if 0 <= idx < len(cache["actors"]):
                selected_actors.append(cache["actors"][idx])

        if not selected_actors:
            await message.answer("Не выбраны адресаты. Напиши номера через запятую.")
            return True

        cache["selected_actors"] = selected_actors

        selected_keys = {a["key"] for a in selected_actors}
        selected_headers = [
            e for e in cache["emails"]
            if (
                _actor_key(e.get("from", "")) in selected_keys
                or _actor_key(e.get("to", "")) in selected_keys
            )
        ]

        selected_headers = _expand_headers_by_threads(
            cache["emails"],
            selected_headers,
        )

        await message.answer(
            f"Выбрано писем: {len(selected_headers)}.\n"
            "Теперь читаю полные письма и вложения по выбранным адресатам..."
        )

        selected_emails = _fetch_selected_full_emails(selected_headers)

        for e in selected_emails:
            e["body"] = (e.get("body") or "")[:3000]
            e["attachments"] = e.get("attachments") or []

        
        # ВАЖНО: сохраняем thread-идентификацию для GPT
        for e in selected_emails:
            e["thread_key"] = f"{_norm_subject(e.get('subject'))}|{_actor_key(e.get('from'))}"

                

        cache["selected"] = selected_emails
        cache["stage"] = "analyzing"

        await message.answer(
            "Адресаты выбраны.\n\n"
            "Запускаю анализ переписки и собираю таблицу акта."
        )

        print("\n========== SECRETARY DEBUG ==========", flush=True)
        print(f"Всего писем: {len(cache['emails'])}", flush=True)
        print(f"Выбрано адресатов: {len(selected_actors)}", flush=True)
        print(f"После фильтра осталось писем: {len(selected_emails)}", flush=True)

        total_chars = sum(len(e.get("subject", "")) for e in selected_emails)
        print(f"Всего символов для GPT: {total_chars}", flush=True)

        import json

        with open("/tmp/raw_selected_emails.json", "w", encoding="utf-8") as f:
            json.dump(
                selected_emails,
                f,
                ensure_ascii=False,
                indent=2,
            )

        print("DEBUG FILE: /tmp/raw_selected_emails.json", flush=True)

       

        cache["selected"] = selected_emails

        await message.answer(
            f"Выбрано писем: {len(selected_emails)}\n"
            "Формирую акт..."
        )

        tasks = _gpt_make_tasks(
            selected_emails,
            cache.get("project", ""),
            cache.get("period_text", "")
        )

        def _fix_status(tasks):
            for t in tasks:
                text = ((t.get("done") or "") + " " + (t.get("topic") or "")).lower()

                if any(x in text for x in [
                    "подготовлен",
                    "направлен",
                    "отправлен",
                    "заключение",
                    "ответ",
                    "проект",
                    "договор",
                    "акт"
                ]):
                    t["status"] = "завершено"
                else:
                    t["status"] = "в работе"

            return tasks


        tasks = _fix_status(tasks)

        if not isinstance(tasks, list):
            await message.answer("Ошибка генерации задач")
            return True

        # 🔥 СТАБИЛИЗАЦИЯ ID ЗАДАЧ (ВАЖНО ДЛЯ РЕДАКТИРОВАНИЯ)
        for i, t in enumerate(tasks):
            t["task_id"] = i + 1

        cache["tasks"] = tasks
        cache["stage"] = "review_tasks"

        await message.answer(
            _tasks_review_text(tasks),
            reply_markup=_tasks_review_keyboard(tasks),
        )

    
    # -----------------------------
    # FALLBACK
    # -----------------------------
    await message.answer(
        "Секретарь активен.\n\n"
        "Команда для старта:\n"
        "составь акт по РГП"
    )
    return True